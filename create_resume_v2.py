import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_margins(cell, top=0, bottom=0, start=0, end=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', start), ('w:right', end)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_horizontal_line(paragraph, color_hex="003399", size=12):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_top_line(paragraph, color_hex="003399", size=12):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(size))
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), color_hex)
    pBdr.append(top)
    pPr.append(pBdr)

def remove_spacing(paragraph):
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1

def create_resume():
    doc = Document()
    
    # A4 size
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    
    # Very narrow margins
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(8.5) # Smaller font to fit one page
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.2

    # Top Header Table
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(14.6)
    table.columns[1].width = Cm(4.0)
    
    cell_info = table.cell(0, 0)
    set_cell_margins(cell_info)
    
    p_name = cell_info.paragraphs[0]
    remove_spacing(p_name)
    run_name = p_name.add_run("陈俊宇  ")
    run_name.font.size = Pt(24)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
    run_en = p_name.add_run("Junyu Chen")
    run_en.font.size = Pt(12)
    run_en.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    p_title = cell_info.add_paragraph()
    remove_spacing(p_title)
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(8)
    r_t = p_title.add_run("大数据技术与工程硕士在读 | 全栈开发 | AI应用与数据工程 | 物联网平台 | 软件设计师")
    r_t.font.bold = True
    r_t.font.size = Pt(9)
    add_horizontal_line(p_title, size=6)
    
    # Add contact info via a nested table to align perfectly
    t_contact = cell_info.add_table(rows=2, cols=3)
    t_contact.autofit = False
    for row in t_contact.rows:
        for c in row.cells:
            set_cell_margins(c)
    t_contact.columns[0].width = Cm(4.5)
    t_contact.columns[1].width = Cm(4.5)
    t_contact.columns[2].width = Cm(5.6)
    
    def set_contact(cell, icon, text):
        p = cell.paragraphs[0]
        remove_spacing(p)
        r_i = p.add_run(icon + " ")
        r_i.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
        p.add_run(text).font.size = Pt(8.5)
        
    set_contact(t_contact.cell(0, 0), "👤", "性别 | 男 | 2002.08")
    set_contact(t_contact.cell(0, 1), "📞", "18983464109")
    set_contact(t_contact.cell(0, 2), "💬", "chenunyu")
    set_contact(t_contact.cell(1, 0), "✉", "1147254646@qq.com")
    set_contact(t_contact.cell(1, 1), "🐧", "1147254646")
    set_contact(t_contact.cell(1, 2), "🌐", "https://junyuc.me")
    
    # Avatar
    cell_avatar = table.cell(0, 1)
    cell_avatar.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_avatar = cell_avatar.paragraphs[0]
    p_avatar.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    remove_spacing(p_avatar)
    avatar_path = r"C:\Users\chenjunyu\Downloads\ChatGPT Image Jun 17, 2026, 07_18_57 PM.png"
    if os.path.exists(avatar_path):
        p_avatar.add_run().add_picture(avatar_path, width=Cm(2.8))
        
    p_gap = doc.add_paragraph()
    remove_spacing(p_gap)
    p_gap.paragraph_format.space_after = Pt(4)
    add_top_line(p_gap, size=6)

    def add_section_header(title, icon="🎯"):
        p = doc.add_paragraph()
        remove_spacing(p)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        
        r_icon = p.add_run(icon + " ")
        r_icon.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
        r_icon.font.size = Pt(12)
        
        r_title = p.add_run(title)
        r_title.font.size = Pt(12)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
        add_horizontal_line(p, size=8)

    # 1. 求职定位
    add_section_header("求职定位", "🎯")
    p = doc.add_paragraph("面向全栈开发、AI应用开发、大数据工程等岗位。具备从需求分析、方案设计、前端后端开发、算法数据处理、接口联调到部署运维的完整项目落地经验，擅长构建高性能、可扩展的系统，做好模型、数据和硬件终端做成可运行、可展示、可维护的系统。")
    remove_spacing(p)
    p.paragraph_format.left_indent = Cm(0.6)

    # 2. 教育背景
    add_section_header("教育背景", "🎓")
    
    def add_edu(school, major, date, bullets):
        p = doc.add_paragraph()
        remove_spacing(p)
        p.paragraph_format.left_indent = Cm(0.6)
        
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(18.0), WD_TAB_ALIGNMENT.RIGHT)
        
        r_bullet = p.add_run("● ")
        r_bullet.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
        
        p.add_run(school).font.bold = True
        p.add_run(" \t" + major)
        p.add_run("\t" + date)
        
        for b in bullets:
            pb = doc.add_paragraph("• " + b)
            remove_spacing(pb)
            pb.paragraph_format.left_indent = Cm(1.0)

    add_edu("长江师范学院", "硕士在读 | 大数据技术与工程", "2028届", [
        "本科阶段 | 交通计算机热力方向",
        "研一期间获一等奖学金；跟随导师参与智能交通预测、大模型工具组网等方向研究。"
    ])
    add_edu("长江师范学院", "本科 | 计算机科学与技术专业", "2025届", [
        "研究方向涉及基于图增强模型的公交乘客上下车站点预测、MQTT 协会工具测控等。",
        "本科毕设设计聚焦城市公交客流OD分析，基于Spark处理公交刷卡与定位数据。"
    ])

    # 3. 核心技能
    add_section_header("核心技能", "🛠")
    
    table_skills = doc.add_table(rows=4, cols=2)
    table_skills.columns[0].width = Cm(2.5)
    table_skills.columns[1].width = Cm(15.5)
    
    skills = [
        ("后端开发", "Java / Spring Boot / MyBatis-Plus / REST API / WebSocket / SSE / Linux / Nginx / Vercel / 服务器部署"),
        ("前端与可视化", "Vue3 / Vite / JavaScript / Ant Design / Vue-数据看板 / 移动端H5 / PC页面开发"),
        ("数据与算法", "Python / Spark / Hadoop / MapReduce / Flink / Kafka / 数据预处理 / OD分析 / 候选召回 / 排序模型 / 模型评估"),
        ("AI与物联网", "大模型工具调用 / MCP / 动态工具剪枝 / ESP32-C3 / MQTT / RFID / 人脸识别 / API / LoRa / Modbus / Netty")
    ]
    
    for i, (k, v) in enumerate(skills):
        cell_k = table_skills.cell(i, 0)
        set_cell_margins(cell_k, top=40, bottom=40)
        p_k = cell_k.paragraphs[0]
        remove_spacing(p_k)
        p_k.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_k = p_k.add_run(k)
        r_k.font.bold = True
        r_k.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "003399")
        cell_k._tc.get_or_add_tcPr().append(shading_elm)
        
        cell_v = table_skills.cell(i, 1)
        set_cell_margins(cell_v, top=40, bottom=40, start=100)
        p_v = cell_v.paragraphs[0]
        remove_spacing(p_v)
        p_v.add_run(v)

    # 4. 项目经历
    add_section_header("项目经历", "💼")
    
    def add_project(name, role, time, bullets):
        p = doc.add_paragraph()
        remove_spacing(p)
        p.paragraph_format.left_indent = Cm(0.6)
        
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(18.0), WD_TAB_ALIGNMENT.RIGHT)
        
        r_bullet = p.add_run("● ")
        r_bullet.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
        
        r_name = p.add_run(name)
        r_name.font.bold = True
        r_name.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
        
        p.add_run(f" ({role}) \t{time}")
        for b in bullets:
            pb = doc.add_paragraph("• " + b)
            remove_spacing(pb)
            pb.paragraph_format.left_indent = Cm(1.0)

    add_project("基于图增强模型的公交下车站点预测系统", "主要负责人 / 全栈开发", "2025.09 - 至今", [
        "面向公交出行场景，结合历史出行记录、线路站点关系、时间与空间特征，构建下车站点候选召回与排序预测流程，重点优化 Top-K 候选中的 Top-1 预测效果。",
        "将预测流程工程化为公交预测工作台，包含综合监控、客流预测、客流调度、算法评估、GIS 管理和客户异常页面。"
    ])
    add_project("智享座位物联网平台 / ESP32 智能座椅系统", "后端开发 / 物联网开发", "2026", [
        "设计实现室内多座位物联网管理看板，硬件端采集温湿度、RFID 卡等座位状态，后端接收 MQTT/HTTP 数据并写入 MySQL。",
        "实现 RFID 与人脸识别认证策略，座位状态控制、实时事件流、前端看板展示以及蒲士报务部署。"
    ])
    add_project("MCP-DTP 大模型 MCP 工具服务器动态工具剪枝方法", "主要负责人", "2026.04 - 至今", [
        "针对工具链增多、模型选择成本高、调用风险增高问题，设计记忆引导的动态工具剪枝方法。",
        "将工具链 active、reserved、blocked 三类管理，规划动态恢复机制与历史记忆引导策略，用于提升工具调用效率与安全性。"
    ])
    add_project("基于微服务架构的农业物联网远程云管理监控系统", "J&K / 后端开发", "2023", [
        "面向智慧农业场景，实现传感器对接数据采集、LoRa/Netty 通信传输、Kafka/Spark 数据解析处理和 Vue 实时监控。",
        "作为队长参与方案设计、后端接口开发、设备数据解析、功能模块和项目文档撰写。"
    ])

    # 5. 荣誉证书
    add_section_header("荣誉证书", "🏆")
    table_awards = doc.add_table(rows=4, cols=2)
    table_awards.columns[0].width = Cm(10.0)
    table_awards.columns[1].width = Cm(8.0)
    awards = [
        ("• 软件设计师 (软考中级) 证书", "• 研究生一年级一等奖学金"),
        ("• 本科阶段国家奖学金", "• 本科阶段优秀毕业生"),
        ("• 2024 年第 17 届中国大学生计算机设计大赛全国二等奖", "• 2023 年第 16 届中国大学生计算机设计大赛重庆市级赛二等奖"),
        ("• 2023 年第十一届全国大学生数字媒体科技作品及创意竞赛国赛一等奖", "• 第七届 iCAN 大学生创新创业大赛重庆赛区二等奖")
    ]
    for i, (a1, a2) in enumerate(awards):
        c1 = table_awards.cell(i, 0)
        p1 = c1.paragraphs[0]
        remove_spacing(p1)
        p1.paragraph_format.left_indent = Cm(0.6)
        p1.add_run(a1)
        
        c2 = table_awards.cell(i, 1)
        p2 = c2.paragraphs[0]
        remove_spacing(p2)
        p2.add_run(a2)

    # 6. 个人优势
    add_section_header("个人优势", "⭐")
    def add_adv(text):
        p = doc.add_paragraph("• " + text)
        remove_spacing(p)
        p.paragraph_format.left_indent = Cm(0.6)
        
    add_adv("学习速度快，能够模块快速掌握新技术并解决复杂工程问题；擅长使用 ChatGPT、Codex、Google、Antigravity 等 AI 辅助开发。")
    add_adv("实践导向强，能把算法、数据、硬件和前后端系统整合成可部署、可演示、可维护的项目。")
    add_adv("具有较好的文档撰写和项目沟通能力，能为复杂项目梳理启动流程、接口路径、排查步骤和部署说明。")

    out_path = r"C:\Users\chenjunyu\Desktop\陈俊宇_复刻简历.docx"
    doc.save(out_path)
    print(f"Resume generated at {out_path}")

if __name__ == "__main__":
    create_resume()
