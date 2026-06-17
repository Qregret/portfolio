import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph, color_hex="003399"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12') # 1.5 pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_resume():
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Header Table
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(1.5)
    
    cell_info = table.cell(0, 0)
    p_name = cell_info.paragraphs[0]
    run_name = p_name.add_run("陈俊宇  ")
    run_name.font.size = Pt(28)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
    run_en = p_name.add_run("Junyu Chen")
    run_en.font.size = Pt(14)
    run_en.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    p_title = cell_info.add_paragraph()
    p_title.add_run("大数据技术与工程硕士在读 | 全栈开发 | AI应用与数据工程 | 物联网平台 | 软件设计师").font.bold = True
    
    p_contact1 = cell_info.add_paragraph()
    p_contact1.add_run("性别 | 男 | 2002.08    18983464109    chenunyu")
    p_contact2 = cell_info.add_paragraph()
    p_contact2.add_run("1147254646@qq.com    1147254646    https://junyuc.me")
    
    # Avatar
    cell_avatar = table.cell(0, 1)
    p_avatar = cell_avatar.paragraphs[0]
    p_avatar.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    avatar_path = r"C:\Users\chenjunyu\Downloads\ChatGPT Image Jun 17, 2026, 07_18_57 PM.png"
    if os.path.exists(avatar_path):
        p_avatar.add_run().add_picture(avatar_path, width=Inches(1.2))
        
    doc.add_paragraph() # spacing

    # Helper function for sections
    def add_section(title):
        p = doc.add_paragraph()
        run = p.add_run("  " + title)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
        add_horizontal_line(p)

    # 1. 求职定位
    add_section("求职定位")
    p = doc.add_paragraph("面向全栈开发、AI应用开发、大数据工程等岗位。具备从需求分析、方案设计、前端后端开发、算法数据处理、接口联调到部署运维的完整项目落地经验，擅长构建高性能、可扩展的系统，做好模型、数据和硬件终端做成可运行、可展示、可维护的系统。")

    # 2. 教育背景
    add_section("教育背景")
    p1 = doc.add_paragraph()
    p1.add_run("长江师范学院        硕士在读 | 大数据技术与工程                                       2028届").font.bold = True
    p1_1 = doc.add_paragraph("• 本科阶段 | 交通计算机热力方向", style='List Bullet')
    p1_2 = doc.add_paragraph("• 研一期间获一等奖学金；跟随导师参与智能交通预测、大模型工具组网等方向研究。", style='List Bullet')

    p2 = doc.add_paragraph()
    p2.add_run("长江师范学院        本科 | 计算机科学与技术专业                                       2024届").font.bold = True
    p2_1 = doc.add_paragraph("• 研究方向涉及基于图增强模型的公交乘客上下车站点预测、MQTT 协会工具测控等。", style='List Bullet')
    p2_2 = doc.add_paragraph("• 本科毕设设计聚焦城市公交客流OD分析，基于Spark处理公交刷卡与定位数据。", style='List Bullet')

    # 3. 核心技能
    add_section("核心技能")
    skills = [
        ("后端开发", "Java / Spring Boot / MyBatis-Plus / REST API / WebSocket / SSE / Linux / Nginx / Vercel / 服务器部署"),
        ("前端与可视化", "Vue3 / Vite / JavaScript / Ant Design / Vue-数据看板 / 移动端H5 / PC页面开发"),
        ("数据与算法", "Python / Spark / Hadoop / MapReduce / Flink / Kafka / 数据预处理 / OD分析 / 候选召回 / 排序模型 / 模型评估"),
        ("AI与物联网", "大模型工具调用 / MCP / 动态工具剪枝 / ESP32-C3 / MQTT / RFID / 人脸识别 / API / LoRa / Modbus / Netty")
    ]
    table_skills = doc.add_table(rows=len(skills), cols=2)
    table_skills.columns[0].width = Inches(1.2)
    table_skills.columns[1].width = Inches(6.0)
    for i, (k, v) in enumerate(skills):
        cell_k = table_skills.cell(i, 0)
        p_k = cell_k.paragraphs[0]
        r_k = p_k.add_run(k)
        r_k.font.bold = True
        r_k.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # hack for background color
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "003399")
        cell_k._tc.get_or_add_tcPr().append(shading_elm)
        
        table_skills.cell(i, 1).text = v

    # 4. 项目经历
    add_section("项目经历")
    
    def add_project(name, role, time, bullets):
        p = doc.add_paragraph()
        run_name = p.add_run(name)
        run_name.font.bold = True
        run_name.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
        p.add_run(f" ({role}) \t\t\t\t\t\t {time}")
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')

    add_project("基于图增强模型的公交下车站点预测系统", "主要负责人 / 全栈开发", "2025.09 - 至今", [
        "面向公交出行场景，结合历史出行记录、线路站点关系、时间与空间特征，构建下车站点候选召回与排序预测流程，重点优化 Top-K 候选中的 Top-1 预测效果。",
        "将预测流程工程化为公交预测工作台，包含综合监控、客流预测、客流调度、算法评估、GIS 管理和客户异常页面。"
    ])
    add_project("智享座位物联网平台 / ESP32 智能座椅系统", "后端开发 / 物联网开发", "2026", [
        "设计实现室内多座位物联网管理看板，硬件端采集温湿度、RFID 卡等座位状态，后端接收 MQTT/HTTP 数据并写入 MySQL。",
        "实现 RFID 与人脸识别认证策略，座位状态控制、实时事件流、前端看板展示以及蒲士报务部署。"
    ])
    add_project("MCP-DTP 大模型 MCP 工具服务器动态工具剪枝方法", "主要负责人", "2026.04 - 至今", [
        "针对工具链增多、模型选择成本高、调用风险增高问题，设计记忆引导的动态工具剪枝方法。",
        "将工具链 active、reserved、blocked 三类管理，规划动态恢复机制与历史记忆引导策略，用于提升工具调用效率与安全性。"
    ])
    add_project("基于微服务架构的农业物联网远程云管理监控系统", "J&K / 后端开发", "2023", [
        "面向智慧农业场景，实现传感器对接数据采集、LoRa/Netty 通信传输、Kafka/Spark 数据解析处理和 Vue 实时监控。",
        "作为队长参与方案设计、后端接口开发、设备数据解析、功能模块和项目文档撰写。"
    ])

    # 5. 荣誉证书
    add_section("荣誉证书")
    table_awards = doc.add_table(rows=4, cols=2)
    awards = [
        ("• 软件设计师 (软考中级) 证书", "• 研究生一年级一等奖学金"),
        ("• 本科阶段国家奖学金", "• 本科阶段优秀毕业生"),
        ("• 2024 年第 17 届中国大学生计算机设计大赛全国二等奖", "• 2023 年第 16 届中国大学生计算机设计大赛重庆市级赛二等奖"),
        ("• 2023 年第十一届全国大学生数字媒体科技作品及创意竞赛国赛一等奖", "• 第七届 iCAN 大学生创新创业大赛重庆赛区二等奖")
    ]
    for i, (a1, a2) in enumerate(awards):
        table_awards.cell(i, 0).text = a1
        table_awards.cell(i, 1).text = a2

    # 6. 个人优势
    add_section("个人优势")
    doc.add_paragraph("学习速度快，能够模块快速掌握新技术并解决复杂工程问题；擅长使用 ChatGPT、Codex、Google、Antigravity 等 AI 辅助开发。", style='List Bullet')
    doc.add_paragraph("实践导向强，能把算法、数据、硬件和前后端系统整合成可部署、可演示、可维护的项目。", style='List Bullet')
    doc.add_paragraph("具有较好的文档撰写和项目沟通能力，能为复杂项目梳理启动流程、接口路径、排查步骤和部署说明。", style='List Bullet')

    out_path = r"C:\Users\chenjunyu\Desktop\陈俊宇_复刻简历.docx"
    doc.save(out_path)
    print(f"Resume generated at {out_path}")

if __name__ == "__main__":
    create_resume()
